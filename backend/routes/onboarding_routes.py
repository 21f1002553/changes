from flask import Blueprint, request, jsonify
from app import db
from utils.file_utils import save_onboarding_file, allowed_file, get_data_root
from models import (
    User, Role, OnBoardingDocument, OnboardingForm, AuditLog, Notification, OfferLetter
)
import os
from flask_jwt_extended import get_jwt_identity, jwt_required
from datetime import datetime
from config import Config
from docx import Document
import json

MAX_FILE_SIZE_MB = 10

onboarding_bp = Blueprint('onboarding', __name__)

# create audit
def create_audit(actor_id, action, target_type, target_id):
    try:
        a1 = AuditLog(actor_id=actor_id, action=action, target_type=target_type, target_id=target_id)
        db.session.add(a1)
        db.session.commit()
    except Exception as e:
        db.session.rollback()


# create notification
def notify(user_id, message):
    try:
        n1 = Notification(recipient_id=user_id, message=message)
        db.session.add(n1)
        db.session.commit()
    except Exception as e:
        db.session.rollback()


#### *Document Management*


# document upload
@onboarding_bp.route('/documents/upload', methods=['POST'])
@jwt_required()
def upload_document():
    f = request.files.get('file')
    if not f:
        return jsonify({'error': 'No file uploaded'}), 400

    employee_id = request.form.get('employee_id')
    if not employee_id:
        return jsonify({'error': 'User ID not provided'}), 400

    user = User.query.get(employee_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404

    current_user_id = get_jwt_identity() # uploader id

    if not allowed_file(f.filename):
        return jsonify({'error': 'File type not allowed'}), 400
    
    doc_type = request.form.get('doc_type')
    if not doc_type:
        return jsonify({'error': 'Document type not provided'}), 400
    try:
        f.stream.seek(0, os.SEEK_END)
        size_mb = f.stream.tell() / (1024 * 1024)
        f.stream.seek(0)
        max_mb = MAX_FILE_SIZE_MB
        if size_mb > max_mb:
            return jsonify({'error': f'File size exceeds the maximum limit of {max_mb} MB'}), 400
    except Exception as e:
        pass

    try:
        path, filename, mime_type = save_onboarding_file(f, employee_id)

    except Exception as e:
        return jsonify({'error':"Failed to save the file", "details":str(e)}), 500


    # create a db row
    try:
        doc = OnBoardingDocument(
            employee_id=str(employee_id),
            doc_type=doc_type,
            filename=filename,
            file_path=path,
            uploaded_by=str(current_user_id) if current_user_id else "None",
            uploaded_at=datetime.utcnow(),
            verified=False
        )
        db.session.add(doc)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({'error':"Failed to persist document","details":str(e)}), 500


    # notify the user
    notify(employee_id, f"onboarding {filename} uploaded successfully")

    # create audit
    create_audit(current_user_id, "upload", "onboarding", doc.id)

    return jsonify({"document": doc.to_dict(), "message": "Document uploaded successfully", "status": "open"}), 201



# get employee documents
@onboarding_bp.route('/documents/<employee_id>', methods=['GET'])
@jwt_required()
def get_documents(employee_id):
    current_user_id = get_jwt_identity()
    try:
        current_user = User.query.get_or_404(current_user_id)
        role = Role.query.filter_by(id=current_user.role_id).first()
        if (current_user.role_id == role.id) or  ((role.name == "admin") or (role.name == "manager") or (role.name == "hr") or (current_user_id == employee_id)):
            documents = OnBoardingDocument.query.filter_by(employee_id=str(employee_id)).all()
            return jsonify([doc.to_dict() for doc in documents]), 200
        else:
            return jsonify({'error': 'Unauthorized role'}), 403
    except Exception as e:
        return jsonify({'error': str(e)}), 500



# verify the document
@onboarding_bp.route('/documents/<doc_id>/verify', methods=['PUT'])
@jwt_required()
def verify_document(doc_id):
    """
    Body JSON:
      { "verifier_id": "<user id>", "verified": true/false }
    Only HR/Admin can verify.
    """
    current_user_id = get_jwt_identity()
    data=request.get_json() or {}
    current_user = User.query.get_or_404(current_user_id)
    role = Role.query.filter_by(id=current_user.role_id).first()
    if (current_user.role_id == role.id) and ((role.name == "admin") or (role.name == "manager") or (role.name == "hr")):
        try:
            doc = OnBoardingDocument.query.get_or_404(doc_id)
            if not doc:
                return jsonify({'error': 'Document not found'}), 404
            if doc.verified:
                return jsonify({'error': 'Document already verified'}), 400
            if doc.verified_by:
                return jsonify({'error': 'Document already verified by another user'}), 400
            doc.verified = True
            doc.verified_by = current_user_id
            doc.verified_at = datetime.now()
            db.session.commit()
            notify(doc.employee_id, f"onboarding {doc.filename} verified successfully")
            create_audit(current_user_id, "verify", "onboarding", doc.id)
            return jsonify({'message': 'Document verified successfully', 'document': doc.to_dict()}), 200
        except Exception as e:
            db.session.rollback()
            return jsonify({'error': str(e)}), 500
    else:
        return jsonify({'error': 'Unauthorized role'}), 403


# onboarding progress employee_id
@onboarding_bp.route('/progress/<employee_id>', methods=['GET'])
@jwt_required()
def onboarding_progress(employee_id):
    current_user_id = get_jwt_identity()
    try:
        current_user = User.query.get_or_404(current_user_id)
        role = Role.query.filter_by(id=current_user.role_id).first()
        required_doc = Config.ONBOARDING_REQUIRED_DOCS
        if role.name not in ["admin", "manager", "hr"] and str(current_user_id) != str(employee_id):
            return jsonify({'error': 'Unauthorized role'}), 403
        
        required_docs = Config.ONBOARDING_REQUIRED_DOCS
        uploaded_docs = OnBoardingDocument.query.filter_by(employee_id=str(employee_id)).all()

        doc_status={}
        completed_docs=0

        for doc_type in required_docs:
            found=[doc for doc in uploaded_docs if doc.doc_type==doc_type]
            present=bool(found)
            verified=any([d.verified for d in found])

            doc_status[doc_type] = {
                "present": present,
                "verified": verified
            }
            
            if present and verified:
                completed_docs+=1
        
        required_forms = list(Config.ONBOARDING_FORM_SCHEMAS.keys())

        form_status={}
        completed_forms_count=0

        for form_type in required_forms:
            last_submission=OnboardingForm.query.filter_by(employee_id=str(employee_id), form_type=form_type).order_by(OnboardingForm.submitted_at.desc()).first()
            completed=bool(last_submission)

            form_status[form_type] = {
                "completed": completed
            }
            
            if completed:
                completed_forms_count+=1
            
        # calculate progress

        total_items=len(required_forms)+len(required_docs)
        completed_items=completed_docs+completed_forms_count
        
        progress_percent=(completed_items/total_items)*100 if total_items>0 else 0

        return jsonify({
            "employee_id": employee_id,
            "progess_percent": progress_percent,
            "doc_status": doc_status,
            "form_status": form_status
        }), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500


#### *Forms & Paperwork*


## generate offer letter
@onboarding_bp.route('/offer-letter/generate', methods=['POST'])
@jwt_required()
def generate_offer_letter():
    """
    body JSON:
    {
        "hr_id": "<HR user ID>",
        "template_vars": {
            "employee_id": "123",
            "candidate_name": "John Doe",
            "job_title": "School Teacher",
            "salary": "5 LPA"
        }
    }
    Only HR/Admin can generate offer letter.
    """
    current_user_id = get_jwt_identity()
    data = request.get_json() or {}
    current_user = User.query.get_or_404(current_user_id)
    role = Role.query.filter_by(id=current_user.role_id).first()
    if (current_user.role_id == role.id) and ((role.name == "admin") or (role.name == "manager") or (role.name == "hr")):
      
        employee_id = data.get('employee_id')
        hr_id = current_user_id
        template_vars = data.get('template_vars',{})
        employee_id = template_vars.get('employee_id')

        if not hr_id or not employee_id:
            return jsonify({"error":"hr_id and employee_id are required"}), 400

        if not template_vars:
            return jsonify({"error":"template_vars are required"}), 400

        if not template_vars.get('employee_id') or not template_vars.get('candidate_name') or not template_vars.get('job_title') or not template_vars.get('salary'):
            return jsonify({"error":"template_vars are required"}), 400
        
        base = get_data_root()
        doc = Document()
        doc.add_heading("Offer Letter", level=1)
        doc.add_paragraph(f"Dear {template_vars.get('candidate_name','')},")
        doc.add_paragraph("")
        doc.add_paragraph("This is an offer letter generated by the HR team.")
        doc.add_paragraph(f"We are pleased to offer you the position of {template_vars.get('job_title','')}.")
        doc.add_paragraph(f"Your salary will be {template_vars.get('salary','')}.")
        doc.add_paragraph("")    
        doc.add_paragraph("We look forward to welcoming you to our team.")
        doc.add_paragraph("")
        doc.add_paragraph("Best regards,")
        doc.add_paragraph("HR Team")
        filename = f"offer_letter_{template_vars.get('employee_id','')}.docx"
        path = os.path.join(base, "uploads", "offer_letters", str(template_vars.get('employee_id','')), filename)
        os.makedirs(os.path.dirname(path), exist_ok=True)

        try:
            doc.save(path)
        except Exception as e:
            return jsonify({"error": "Failed to save offer letter", "detail": str(e)}), 500
        
        try:
            record = OfferLetter(
                employee_id=str(employee_id),
                created_by=str(hr_id),
                status='sent',
                file_path=path,
                filename=filename
            )
            db.session.add(record)
            db.session.commit()


        except Exception as e:
            db.session.rollback()
            return jsonify({"error": "Failed to save offer letter", "detail": str(e)}), 500
    
        try:
            doc_record = OnBoardingDocument(
                employee_id=str(employee_id),
                uploaded_by=str(hr_id),
                doc_type="offer_letter",
                filename=filename,
                file_path=path
            )
            db.session.add(doc_record)
            db.session.commit()

        except Exception as e:
            db.session.rollback()
            record.status = 'failed'
            db.session.commit()
            return jsonify({"error": "Failed to save offer letter", "detail": str(e)}), 500
        
        notify(employee_id, "Your offer letter has been generated")
        create_audit(hr_id, "generate", "offer letter", record.id)
        
        return jsonify({"offer": record.to_dict(), "document": doc_record.to_dict()}), 201
    
    else:
        return jsonify({'error': 'Unauthorized role'}), 403
            
# get form based on form type

@onboarding_bp.route('/forms/<form_type>', methods=['GET'])
def forms(form_type):
    schemas=getattr(Config, "ONBOARDING_FORM_SCHEMAS", {}) or {}
    schema=schemas.get(form_type)
    if not schema:
        return jsonify({'error': 'Form not found'}), 404
    
    employee_id=request.args.get('employee_id')
    prefilled={}
    if employee_id:
        last=OnboardingForm.query.filter_by(employee_id=employee_id, form_type=form_type).order_by(OnboardingForm.submitted_at.desc()).first()
        if last:
            # payload stored as JSON string in DB; parse safely
            try:
                prefilled = last.get_payload() if hasattr(last, "get_payload") else json.loads(last.payload)
            except Exception:
                prefilled = {}


    return jsonify({'form': form_type, 'schema': schema, 'prefilled': prefilled}), 200


# submit form
@onboarding_bp.route('/forms/submit',methods=['POST'])
@jwt_required()
def submit_form():
    """
    Accepts either:
     - JSON body: { employee_id, form_type, form_data }
     - multipart/form-data: file, employee_id, form_type, form_data (form_data MUST be JSON string)

    Behavior:
     - Validate required fields per Config.ONBOARDING_FORM_SCHEMAS
     - Create an OnboardingForm row (model name: OnboardingForm)
     - If file uploaded (and allowed by schema), save it via save_onboarding_file() and create OnBoardingDocument
     - Notify and audit
    """

    current_user_id = get_jwt_identity()
    if not current_user_id:
        return jsonify({'error': 'Unauthorized'}), 401

    is_multipart = request.content_type.startswith('multipart/form-data')
    file=None
    try:
        if is_multipart:
            file=request.files.get('file')
            employee_id=request.form.get('employee_id')
            form_type=request.form.get('form_type')
            form_data=request.form.get('form_data') or "{}"
            try:
                form_data=json.loads(form_data)
            except Exception as e:
                return jsonify({'error': 'Invalid form data'}), 400
        else:
            payload=request.get_json() or {}
            employee_id=payload.get('employee_id')
            form_type=payload.get('form_type')
            form_data=payload.get('form_data')
    except Exception as e:
        return jsonify({"error": "Invalid request", "detail":str(e)}), 400
    

    if not employee_id or not form_type or not form_data:
        return jsonify({'error': 'Missing required fields'}), 400

    schemas=getattr(Config, "ONBOARDING_FORM_SCHEMAS", {}) or {}
    schema=schemas.get(form_type)
    if not schema:
        return jsonify({'error': 'Form not found'}), 404
    
    missing=[]
    for field in schema.get("fields",[]):
        if field.get("required"):
            name=field.get("name")
            val=form_data.get(name)

            if val is None or (isinstance(val, str) and val.strip() == "") or (isinstance(val, list) and len(val) == 0):
                missing.append(name)

    if missing:
        return jsonify({'error': f'Missing required fields: {", ".join(missing)}'}), 400
    
    try:
        # save form submission
        submission=OnboardingForm(
            employee_id=str(employee_id),
            form_type=form_type,
            payload=json.dumps(form_data),
            submitted_by=str(current_user_id),
            submitted_at=datetime.utcnow()
        )
        db.session.add(submission)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": "Failed to save form", "detail": str(e)}), 500
    
    signed_doc=None
    if is_multipart and file:
        allow_signed=schema.get("allow_upload_signed", False)
        if not allow_signed:
            return jsonify({'error': 'Signed file upload not allowed for this form'}), 400

        try:
            file.stream.seek(0, os.SEEK_END)
            size_mb = file.stream.tell() / (1024 * 1024)
            file.stream.seek(0)
            max_mb = MAX_FILE_SIZE_MB
            if size_mb > max_mb:
                return jsonify({'error': f'File size exceeds the maximum limit of {max_mb} MB'}), 400
        except Exception as e:
            pass
        
        ext=allowed_file(file.filename)
        if not ext:
            return jsonify({'error': 'File type not allowed'}), 400

        try:
            path, filename, mime_type = save_onboarding_file(file, employee_id)

            signed_doc=OnBoardingDocument(
                employee_id=str(employee_id),
                form_type=form_type,
                filename=filename,
                file_path=path,
                created_at=datetime.utcnow(),
                verified=True,
                uploaded_by=str(current_user_id),
            )
            db.session.add(signed_doc)
            db.session.commit()

        except Exception as e:
            signed_doc=None
            return jsonify({"error": "Failed to save file", "detail": str(e)}), 500
        
        try:
            submission.signed_doc_id = signed_doc.id
            # optionally put pointer in payload as well
            payload_obj = submission.get_payload()
            payload_obj['signed_doc_id'] = signed_doc.id
            submission.payload = json.dumps(payload_obj)
            db.session.commit()
        except Exception:
            db.session.rollback()
            pass

    try:
        notify(employee_id, f"Your {form_type} form has been signed and submitted.")
        create_audit(current_user_id, "submit", "onboarding", submission.id)
    except Exception as e:
        return jsonify({"error": "Failed to notify", "detail": str(e)}), 500
    
    response = {"submission": submission.to_dict(), "signed_doc": signed_doc.to_dict() if signed_doc else None}

    return jsonify(response), 201