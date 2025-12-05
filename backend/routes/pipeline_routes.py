from flask import Blueprint, request, jsonify
from app import db
from models import (
    HiringPipelineStage, CandidatePipelineStage, InterviewerAssignment, 
    Scorecard, User, JobPost, Application, OfferLetter, Notification
)
from datetime import datetime
from flask_jwt_extended import jwt_required, get_jwt_identity
from docx import Document
from utils.file_utils import get_data_root
import os

pipeline_bp = Blueprint('pipeline', __name__)

@pipeline_bp.route('/stages', methods=['GET'])
def get_pipeline_stages():
    """Get all hiring pipeline stages"""
    try:
        stages = HiringPipelineStage.query.filter_by(is_active=True).order_by(HiringPipelineStage.order_index).all()
        
        # Get stage statistics
        stages_data = []
        for stage in stages:
            stage_data = stage.to_dict()
            
            # Count candidates in this stage
            candidate_count = CandidatePipelineStage.query.filter_by(
                stage_id=stage.id,
                status='active'
            ).count()
            
            stage_data['candidate_count'] = candidate_count
            stages_data.append(stage_data)
        
        return jsonify({
            'stages': stages_data,
            'total_stages': len(stages_data)
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@pipeline_bp.route('/stages', methods=['POST'])
def create_pipeline_stage():
    """Create a new pipeline stage"""
    try:
        data = request.get_json()
        
        if not data.get('name'):
            return jsonify({'error': 'Stage name is required'}), 400
        
        # Get next order index
        max_order = db.session.query(db.func.max(HiringPipelineStage.order_index)).scalar() or 0
        
        stage = HiringPipelineStage(
            name=data['name'],
            description=data.get('description', ''),
            order_index=data.get('order_index', max_order + 1),
            is_active=data.get('is_active', True)
        )
        
        db.session.add(stage)
        db.session.commit()
        
        return jsonify({
            'message': 'Pipeline stage created successfully',
            'stage': stage.to_dict()
        }), 201
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@pipeline_bp.route('/candidates/<stage_id>', methods=['GET'])
def get_candidates_in_stage(stage_id):
    """Get all candidates in a specific pipeline stage"""
    try:
        # Validate stage exists
        stage = HiringPipelineStage.query.get_or_404(stage_id)
        
        # Get query parameters
        job_id = request.args.get('job_id')
        limit = request.args.get('limit', 50, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        # Build query
        query = CandidatePipelineStage.query.filter_by(
            stage_id=stage_id,
            status='active'
        )
        
        if job_id:
            query = query.filter_by(job_id=job_id)
        
        # Order by most recently entered
        query = query.order_by(CandidatePipelineStage.entered_at.desc())
        
        # Apply pagination
        candidate_stages = query.offset(offset).limit(limit).all()
        total_count = query.count()
        
        # Format response with additional candidate data
        candidates_data = []
        for candidate_stage in candidate_stages:
            candidate_data = candidate_stage.to_dict()
            
            # Add candidate details
            if candidate_stage.candidate:
                candidate_data['candidate_details'] = {
                    'id': candidate_stage.candidate.id,
                    'name': candidate_stage.candidate.name,
                    'email': candidate_stage.candidate.email
                }
            
            # Add application details
            if candidate_stage.job_id:
                application = Application.query.filter_by(
                    candidate_id=candidate_stage.candidate_id,
                    job_id=candidate_stage.job_id
                ).first()
                
                if application:
                    candidate_data['application_details'] = {
                        'id': application.id,
                        'status': application.status,
                        'score': application.score,
                        'applied_at': application.applied_at.isoformat() if application.applied_at else None
                    }
            
            candidates_data.append(candidate_data)
        
        return jsonify({
            'stage': stage.to_dict(),
            'candidates': candidates_data,
            'pagination': {
                'total': total_count,
                'limit': limit,
                'offset': offset,
                'has_more': offset + limit < total_count
            }
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@pipeline_bp.route('/candidates', methods=['GET'])
def get_all_pipeline_candidates():
    """Get all candidates across all pipeline stages"""
    try:
        job_id = request.args.get('job_id')
        limit = request.args.get('limit', 100, type=int)
        offset = request.args.get('offset', 0, type=int)
        

        query = CandidatePipelineStage.query.filter_by(status='active')
        
        if job_id:
            query = query.filter_by(job_id=job_id)
        

        query = query.join(HiringPipelineStage).order_by(
            HiringPipelineStage.order_index,
            CandidatePipelineStage.entered_at.desc()
        )
        
        candidate_stages = query.offset(offset).limit(limit).all()
        

        candidates_data = []
        for candidate_stage in candidate_stages:
            candidate_data = candidate_stage.to_dict()
            

            if candidate_stage.candidate:
                candidate_data['candidate_details'] = {
                    'id': candidate_stage.candidate.id,
                    'name': candidate_stage.candidate.name,
                    'email': candidate_stage.candidate.email
                }
            

            if candidate_stage.job_id:
                application = Application.query.filter_by(
                    candidate_id=candidate_stage.candidate_id,
                    job_id=candidate_stage.job_id
                ).first()
                
                if application:
                    candidate_data['application_details'] = {
                        'id': application.id,
                        'job_title': application.job.title if application.job else 'Unknown',
                        'status': application.status,
                        'score': application.score,
                        'applied_at': application.applied_at.isoformat() if application.applied_at else None
                    }
            

            candidate_data['stage_info'] = {
                'id': candidate_stage.stage.id,
                'name': candidate_stage.stage.name,
                'order_index': candidate_stage.stage.order_index
            }
            
            candidates_data.append(candidate_data)
        
        return jsonify({
            'candidates': candidates_data,
            'total': query.count()
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@pipeline_bp.route('/candidates/<candidate_id>/move', methods=['PUT'])
def move_candidate_in_pipeline(candidate_id):
    """Move a candidate to a different pipeline stage"""
    try:
        data = request.get_json()
        
        new_stage_id = data.get('stage_id')
        job_id = data.get('job_id')
        notes = data.get('notes', '')
        moved_by_id = data.get('moved_by_id')
        
        if not new_stage_id:
            return jsonify({'error': 'stage_id is required'}), 400
        
        if not job_id:
            return jsonify({'error': 'job_id is required'}), 400
        
        # Validate entities exist
        candidate = User.query.get_or_404(candidate_id)
        new_stage = HiringPipelineStage.query.get_or_404(new_stage_id)
        job = JobPost.query.get_or_404(job_id)
        
        # Find current active stage for this candidate and job
        current_stage = CandidatePipelineStage.query.filter_by(
            candidate_id=candidate_id,
            job_id=job_id,
            status='active'
        ).first()
        
        # Complete current stage if exists
        if current_stage:
            current_stage.status = 'completed'
            current_stage.completed_at = datetime.utcnow()
        
        # Create new stage entry
        new_candidate_stage = CandidatePipelineStage(
            candidate_id=candidate_id,
            job_id=job_id,
            stage_id=new_stage_id,
            notes=notes,
            moved_by_id=moved_by_id,
            status='active'
        )
        
        db.session.add(new_candidate_stage)
        db.session.commit()
        
        return jsonify({
            'message': f'Candidate moved to {new_stage.name} stage',
            'candidate_id': candidate_id,
            'candidate_name': candidate.name,
            'job_id': job_id,
            'job_title': job.title,
            'previous_stage': current_stage.stage.name if current_stage else None,
            'new_stage': new_stage.name,
            'stage_entry': new_candidate_stage.to_dict()
        }), 200
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@pipeline_bp.route('/assign-interviewer', methods=['POST'])
def assign_interviewer():
    """Assign an interviewer to a candidate"""
    try:
        data = request.get_json()
        
        required_fields = ['candidate_id', 'job_id', 'interviewer_id']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
        
        candidate_id = data['candidate_id']
        job_id = data['job_id']
        interviewer_id = data['interviewer_id']

        # Validate entities exist
        candidate = User.query.get_or_404(candidate_id)
        job = JobPost.query.get_or_404(job_id)
        interviewer = User.query.get_or_404(interviewer_id)
        
        # Parse scheduled time if provided
        scheduled_at = None
        if data.get('scheduled_at'):
            try:
                scheduled_at = datetime.fromisoformat(data['scheduled_at'].replace('Z', '+00:00'))
            except ValueError:
                return jsonify({'error': 'Invalid scheduled_at format. Use ISO 8601 format'}), 400
        
        # Create interviewer assignment
        assignment = InterviewerAssignment(
            candidate_id=candidate_id,
            job_id=job_id,
            interviewer_id=interviewer_id,
            interview_type=data.get('interview_type', 'technical'),
            scheduled_at=scheduled_at,
            duration_minutes=data.get('duration_minutes', 60),
            meeting_link=data.get('meeting_link', ''),
            assigned_by_id=data.get('assigned_by_id')
        )
        
        db.session.add(assignment)
        db.session.commit()
        
        return jsonify({
            'message': 'Interviewer assigned successfully',
            'assignment': assignment.to_dict()
        }), 201
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@pipeline_bp.route('/scorecard', methods=['POST'])
def create_scorecard():
    """Create an interview scorecard"""
    try:
        data = request.get_json()
        
        required_fields = ['interview_assignment_id', 'interviewer_id']
        for field in required_fields:
            if not data.get(field):
                return jsonify({'error': f'{field} is required'}), 400
        
        # Get the interview assignment
        assignment = InterviewerAssignment.query.get_or_404(data['interview_assignment_id'])
        
        # Validate interviewer matches
        if assignment.interviewer_id != data['interviewer_id']:
            return jsonify({'error': 'Interviewer ID does not match assignment'}), 400
        
        # Check if scorecard already exists
        existing_scorecard = Scorecard.query.filter_by(
            interview_assignment_id=data['interview_assignment_id'],
            interviewer_id=data['interviewer_id']
        ).first()
        
        if existing_scorecard and not data.get('force_update', False):
            return jsonify({
                'error': 'Scorecard already exists for this interview. Use force_update=true to update.'
            }), 409
        
        if existing_scorecard:
            # Update existing scorecard
            scorecard = existing_scorecard
        else:
            # Create new scorecard
            scorecard = Scorecard(
                interview_assignment_id=data['interview_assignment_id'],
                interviewer_id=data['interviewer_id'],
                candidate_id=assignment.candidate_id,
                job_id=assignment.job_id
            )
        
        # Update scorecard fields
        scorecard.technical_score = data.get('technical_score')
        scorecard.communication_score = data.get('communication_score')
        scorecard.problem_solving_score = data.get('problem_solving_score')
        scorecard.cultural_fit_score = data.get('cultural_fit_score')
        scorecard.strengths = data.get('strengths', '')
        scorecard.weaknesses = data.get('weaknesses', '')
        scorecard.feedback_notes = data.get('feedback_notes', '')
        scorecard.recommendation = data.get('recommendation', '')
        scorecard.is_final = data.get('is_final', False)
        
        # Calculate overall score if individual scores are provided
        scores = [
            scorecard.technical_score,
            scorecard.communication_score,
            scorecard.problem_solving_score,
            scorecard.cultural_fit_score
        ]
        
        valid_scores = [s for s in scores if s is not None]
        if valid_scores:
            scorecard.overall_score = sum(valid_scores) / len(valid_scores)
        elif data.get('overall_score') is not None:
            scorecard.overall_score = data['overall_score']
        
        if not existing_scorecard:
            db.session.add(scorecard)
        
        # Update interview assignment status
        assignment.status = 'completed'
        
        db.session.commit()
        
        return jsonify({
            'message': 'Scorecard created successfully',
            'scorecard': scorecard.to_dict(),
            'interview_assignment': assignment.to_dict()
        }), 201
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500

@pipeline_bp.route('/scorecards', methods=['GET'])
def get_scorecards():
    """Get scorecards with optional filtering"""
    try:
        # Get query parameters
        candidate_id = request.args.get('candidate_id')
        job_id = request.args.get('job_id')
        interviewer_id = request.args.get('interviewer_id')
        min_score = request.args.get('min_score', type=float)
        recommendation = request.args.get('recommendation')
        limit = request.args.get('limit', 20, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        # Build query
        query = Scorecard.query
        
        if candidate_id:
            query = query.filter(Scorecard.candidate_id == candidate_id)
        if job_id:
            query = query.filter(Scorecard.job_id == job_id)
        if interviewer_id:
            query = query.filter(Scorecard.interviewer_id == interviewer_id)
        if min_score is not None:
            query = query.filter(Scorecard.overall_score >= min_score)
        if recommendation:
            query = query.filter(Scorecard.recommendation == recommendation)
        
        # Order by most recent
        query = query.order_by(Scorecard.submitted_at.desc())
        
        # Apply pagination
        scorecards = query.offset(offset).limit(limit).all()
        total_count = query.count()
        
        return jsonify({
            'scorecards': [scorecard.to_dict() for scorecard in scorecards],
            'pagination': {
                'total': total_count,
                'limit': limit,
                'offset': offset,
                'has_more': offset + limit < total_count
            }
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@pipeline_bp.route('/interview-assignments', methods=['GET'])
def get_interview_assignments():
    """Get interview assignments with optional filtering"""
    try:
        # Get query parameters
        candidate_id = request.args.get('candidate_id')
        job_id = request.args.get('job_id')
        interviewer_id = request.args.get('interviewer_id')
        status = request.args.get('status')
        upcoming_only = request.args.get('upcoming_only', 'false').lower() == 'true'
        limit = request.args.get('limit', 20, type=int)
        offset = request.args.get('offset', 0, type=int)
        
        # Build query
        query = InterviewerAssignment.query
        
        if candidate_id:
            query = query.filter(InterviewerAssignment.candidate_id == candidate_id)
        if job_id:
            query = query.filter(InterviewerAssignment.job_id == job_id)
        if interviewer_id:
            query = query.filter(InterviewerAssignment.interviewer_id == interviewer_id)
        if status:
            query = query.filter(InterviewerAssignment.status == status)
        if upcoming_only:
            query = query.filter(
                InterviewerAssignment.scheduled_at > datetime.utcnow(),
                InterviewerAssignment.status == 'scheduled'
            )
        
        # Order by scheduled time
        query = query.order_by(InterviewerAssignment.scheduled_at.desc())
        
        # Apply pagination
        assignments = query.offset(offset).limit(limit).all()
        total_count = query.count()
        
        return jsonify({
            'interview_assignments': [assignment.to_dict() for assignment in assignments],
            'pagination': {
                'total': total_count,
                'limit': limit,
                'offset': offset,
                'has_more': offset + limit < total_count
            }
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@pipeline_bp.route('/stats', methods=['GET'])
def get_pipeline_stats():
    """Get pipeline statistics"""
    try:
        job_id = request.args.get('job_id')
        
        # Get stage statistics
        stages = HiringPipelineStage.query.filter_by(is_active=True).order_by(HiringPipelineStage.order_index).all()
        
        stage_stats = []
        total_candidates = 0
        
        for stage in stages:
            query = CandidatePipelineStage.query.filter_by(
                stage_id=stage.id,
                status='active'
            )
            
            if job_id:
                query = query.filter_by(job_id=job_id)
            
            candidate_count = query.count()
            total_candidates += candidate_count
            
            stage_stats.append({
                'stage': stage.to_dict(),
                'candidate_count': candidate_count
            })
        
        # Get interview statistics
        interview_query = InterviewerAssignment.query
        if job_id:
            interview_query = interview_query.filter_by(job_id=job_id)
        
        interview_stats = {
            'total_interviews': interview_query.count(),
            'scheduled': interview_query.filter_by(status='scheduled').count(),
            'completed': interview_query.filter_by(status='completed').count(),
            'cancelled': interview_query.filter_by(status='cancelled').count()
        }
        
        # Get scorecard statistics
        scorecard_query = Scorecard.query
        if job_id:
            scorecard_query = scorecard_query.filter_by(job_id=job_id)
        
        scorecard_stats = {
            'total_scorecards': scorecard_query.count(),
            'hire_recommendations': scorecard_query.filter_by(recommendation='hire').count(),
            'strong_hire_recommendations': scorecard_query.filter_by(recommendation='strong_hire').count(),
            'no_hire_recommendations': scorecard_query.filter_by(recommendation='no_hire').count(),
            'average_score': scorecard_query.with_entities(db.func.avg(Scorecard.overall_score)).scalar() or 0
        }
        
        return jsonify({
            'job_id': job_id,
            'pipeline_stats': {
                'total_candidates': total_candidates,
                'stages': stage_stats
            },
            'interview_stats': interview_stats,
            'scorecard_stats': scorecard_stats
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# Offer Letter Routes

@pipeline_bp.route('/send-offer', methods=['POST'])
@jwt_required()
def send_offer_letter():
    """Send offer letter to candidate in Offer stage"""
    try:
        current_user_id = get_jwt_identity()
        data = request.get_json()
        
        candidate_id = data.get('candidate_id')
        job_id = data.get('job_id')
        salary = data.get('salary')
        
        if not all([candidate_id, job_id, salary]):
            return jsonify({'error': 'candidate_id, job_id, and salary are required'}), 400
        
        # Validate entities
        candidate = User.query.get_or_404(candidate_id)
        job = JobPost.query.get_or_404(job_id)
        
        # Check if candidate is in Offer stage
        offer_stage = HiringPipelineStage.query.filter_by(name='Offer').first()
        if not offer_stage:
            return jsonify({'error': 'Offer stage not found in pipeline'}), 404
        
        candidate_stage = CandidatePipelineStage.query.filter_by(
            candidate_id=candidate_id,
            job_id=job_id,
            stage_id=offer_stage.id,
            status='active'
        ).first()
        
        if not candidate_stage:
            return jsonify({'error': 'Candidate is not in Offer stage'}), 400
        
        # Check if offer already exists
        existing_offer = OfferLetter.query.filter_by(
            employee_id=candidate_id,
            job_id=job_id
        ).first()
        
        if existing_offer and existing_offer.acceptance_status != 'rejected':
            return jsonify({
                'error': 'Offer letter already sent to this candidate',
                'offer': existing_offer.to_dict()
            }), 409
        
        # Generate offer letter document
        base = get_data_root()
        doc = Document()
        doc.add_heading("Offer Letter", level=1)
        doc.add_paragraph(f"Dear {candidate.name},")
        doc.add_paragraph("")
        doc.add_paragraph("We are pleased to offer you the position of:")
        doc.add_paragraph(f"Position: {job.title}")
        doc.add_paragraph(f"Salary: {salary}")
        doc.add_paragraph(f"Location: {job.location}")
        doc.add_paragraph("")
        doc.add_paragraph("We believe you will be a great addition to our team and look forward to working with you.")
        doc.add_paragraph("")
        doc.add_paragraph("Please accept this offer through the candidate portal.")
        doc.add_paragraph("")
        doc.add_paragraph("Best regards,")
        doc.add_paragraph("HR Team")
        
        filename = f"offer_letter_{candidate_id}_{job_id}.docx"
        path = os.path.join(base, "uploads", "offer_letters", str(candidate_id), filename)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        
        try:
            doc.save(path)
        except Exception as e:
            return jsonify({'error': 'Failed to save offer letter', 'detail': str(e)}), 500
        
        # Create offer letter record
        offer = OfferLetter(
            employee_id=candidate_id,
            job_id=job_id,
            created_by=current_user_id,
            status='sent',
            acceptance_status='pending',
            filename=filename,
            file_path=path,
            notes=data.get('notes', '')
        )
        
        db.session.add(offer)
        
        # Create notification
        notification = Notification(
            recipient_id=candidate_id,
            message=f"You have received an offer letter for the position of {job.title}. Please review and accept."
        )
        db.session.add(notification)
        
        db.session.commit()
        
        return jsonify({
            'message': 'Offer letter sent successfully',
            'offer': offer.to_dict()
        }), 201
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@pipeline_bp.route('/accept-offer/<offer_id>', methods=['PUT'])
@jwt_required()
def accept_offer_letter(offer_id):
    """Candidate accepts or rejects offer letter"""
    try:
        current_user_id = get_jwt_identity()
        data = request.get_json()
        
        acceptance = data.get('acceptance')  # 'accepted' or 'rejected'
        
        if acceptance not in ['accepted', 'rejected']:
            return jsonify({'error': 'acceptance must be "accepted" or "rejected"'}), 400
        
        # Get offer letter
        offer = OfferLetter.query.get_or_404(offer_id)
        
        # Verify candidate
        if offer.employee_id != current_user_id:
            return jsonify({'error': 'Unauthorized to accept this offer'}), 403
        
        # Check if already processed
        if offer.acceptance_status != 'pending':
            return jsonify({
                'error': f'Offer already {offer.acceptance_status}',
                'offer': offer.to_dict()
            }), 400
        
        # Update offer
        offer.acceptance_status = acceptance
        offer.accepted_at = datetime.utcnow()
        offer.status = acceptance
        
        # Create notification for HR
        hr_users = User.query.join(User.role).filter(
            db.or_(
                db.text("roles.name = 'hr'"),
                db.text("roles.name = 'admin'")
            )
        ).all()
        
        for hr_user in hr_users:
            notification = Notification(
                recipient_id=hr_user.id,
                message=f"{offer.employee.name} has {acceptance} the offer for {offer.job.title}"
            )
            db.session.add(notification)
        
        db.session.commit()
        
        return jsonify({
            'message': f'Offer {acceptance} successfully',
            'offer': offer.to_dict()
        }), 200
    
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@pipeline_bp.route('/offer-status/<candidate_id>/<job_id>', methods=['GET'])
@jwt_required()
def get_offer_status(candidate_id, job_id):
    """Get offer letter status for a candidate and job"""
    try:
        offer = OfferLetter.query.filter_by(
            employee_id=candidate_id,
            job_id=job_id
        ).order_by(OfferLetter.created_at.desc()).first()
        
        if not offer:
            return jsonify({
                'has_offer': False,
                'offer': None
            }), 200
        
        return jsonify({
            'has_offer': True,
            'offer': offer.to_dict()
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@pipeline_bp.route('/my-offers', methods=['GET'])
@jwt_required()
def get_my_offers():
    """Get all offer letters for current user (candidate)"""
    try:
        current_user_id = get_jwt_identity()
        
        offers = OfferLetter.query.filter_by(
            employee_id=current_user_id
        ).order_by(OfferLetter.created_at.desc()).all()
        
        return jsonify({
            'offers': [offer.to_dict() for offer in offers],
            'total': len(offers)
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500